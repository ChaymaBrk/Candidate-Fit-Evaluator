import pdfplumber
from docx import Document
from fastapi import UploadFile
import io
import re
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for parsing PDF and DOCX documents"""
    
    @staticmethod
    async def parse_pdf(file: UploadFile) -> str:
        """Parse PDF file and extract text"""
        try:
            content = await file.read()
            
            text = ""
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return DocumentParser._clean_text(text)
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    async def parse_docx(file: UploadFile) -> str:
        """Parse DOCX file and extract text"""
        try:
            content = await file.read()
            doc = Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return DocumentParser._clean_text(text)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    async def parse_txt(file: UploadFile) -> str:
        """Parse plain text file"""
        try:
            content = await file.read()
            text = content.decode('utf-8')
            return DocumentParser._clean_text(text)
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            raise Exception(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', '', text)
        
        # Normalize line breaks
        text = text.replace('\n', ' ').replace('\r', ' ')
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    async def parse_document(file: UploadFile) -> str:
        """Parse document based on file extension"""
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            return await DocumentParser.parse_pdf(file)
        elif filename.endswith('.docx'):
            return await DocumentParser.parse_docx(file)
        elif filename.endswith('.txt'):
            return await DocumentParser.parse_txt(file)
        else:
            raise Exception(f"Unsupported file type: {filename}")

# Test runner
if __name__ == "__main__":
    import asyncio
    logging.basicConfig(level=logging.INFO)
    
    async def test():
        # This part of the test runner needs to be updated to use UploadFile
        # For now, it's commented out as the original code used Path and open
        # If you have a file path, you can uncomment and replace with your path
        # sample_path = Path(r"C:\Users\Asus\OneDrive\Bureau\Test\ChaymaBarkaouiCV.pdf")  # Change to your test file
        # with sample_path.open("rb") as f:
        #     file = UploadFile(filename=sample_path.name, file=f)
        #     print(await DocumentParser.parse_document(file))
        print("Test runner is not fully functional with current setup.")
    
    asyncio.run(test())